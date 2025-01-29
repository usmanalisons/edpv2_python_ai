# app/utils/file_loader.py
import fitz
from docx import Document
import pytesseract
from PIL import Image
import io
import logging
import re
from app.core.config import settings
# from autocorrect import Speller
# from langdetect import detect, DetectorFactory

# DetectorFactory.seed = 0 
# spell = Speller(lang='en')
pytesseract.pytesseract.tesseract_cmd = settings.PYTESSERACT_LOCAL_PATH

def format_file_text(text, apply_spell_check=False):
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'\.{3,}', '', text)
    text = re.sub(r'\bcontrolled\b', '', text, flags=re.IGNORECASE)
    # try:
    #     if detect(text) != 'en':
    #         return ""
    # except Exception as e:
    #     logging.error(f"Language detection failed: {e}")
    #     return ""

    return text.strip()

def load_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    text_pages = []
    for page_num, page in enumerate(doc):
        try:
            text = format_file_text(page.get_text(), apply_spell_check=False)
            if not text.strip() or len(text) < 50:
                logging.info(f"Performing OCR on page {page_num+1} of {file_path}.")
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                ocr_text = pytesseract.image_to_string(img)
                text = format_file_text(ocr_text, apply_spell_check=True)
            
            if not text.strip() or len(text) > 50:
                text_pages.append({
                    "page_number": page_num + 1,
                    "text": text,
                })
        except Exception as e:
            logging.error(f"Error extracting text from page {page_num+1}: {e}")
    return text_pages


# def load_text_from_pdf(file_path):
#     doc = fitz.open(file_path)
#     text_pages = []
#     for page_num, page in enumerate(doc):
#         try:
#             images = page.get_images(full=True)
#             image_found = len(images) > 0
#             text = format_file_text(page.get_text(), apply_spell_check=False)
#             if image_found:
#                 logging.info(f"Performing OCR on page {page_num + 1} of {file_path}.")
#                 pix = page.get_pixmap()
#                 img_data = pix.tobytes("png")
#                 img = Image.open(io.BytesIO(img_data))
#                 ocr_text = pytesseract.image_to_string(img)
#                 text = format_file_text(ocr_text, apply_spell_check=True)

#             if text.strip() and len(text) > 100:
#                 text_pages.append({
#                     "page_number": page_num + 1,
#                     "text": text,
#                 })
        
#         except Exception as e:
#             logging.error(f"Error extracting text from page {page_num + 1}: {e}")
    
#     return text_pages

def load_text_from_docx(file_path): 
    doc = Document(file_path)
    text_pages = []
    for page_num, para in enumerate(doc.paragraphs):
        text_pages.append({
            "page_number": page_num + 1,
            "text": para.text.strip(),
        })
    return text_pages

def load_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return [{
        "page_number": 1,
        "text": format_file_text(text),
    }]

def load_text_from_file(file_path):
    try:
        file_path = file_path.lower()
        if file_path.endswith('.pdf'):
            return load_text_from_pdf(file_path)
        elif file_path.endswith('.docx') or file_path.endswith('.doc'):
            return load_text_from_docx(file_path)
        elif file_path.endswith('.txt'):
            return load_text_from_txt(file_path)
        else:
            logging.info(f'Unsupported file type: {file_path}')
            return []
    except Exception as e:
        logging.info(f'Invalid file: {file_path}')
        return []



# import fitz
# from docx import Document
# import pytesseract
# from PIL import Image, ImageEnhance
# import io
# import logging
# import re
# from app.core.config import settings
# from autocorrect import Speller
# from langdetect import detect, DetectorFactory

# # Configure logging
# logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# # Initialize language tools
# DetectorFactory.seed = 0
# spell = Speller(lang='en')
# pytesseract.pytesseract.tesseract_cmd = settings.PYTESSERACT_LOCAL_PATH

# def preprocess_image(img):
#     """Preprocess image for better OCR results."""
#     img = img.convert("L")  # Convert to grayscale
#     img = img.resize((img.width * 2, img.height * 2), Image.Resampling.LANCZOS)  # Upscale image
#     enhancer = ImageEnhance.Contrast(img)
#     img = enhancer.enhance(2)  # Increase contrast
#     img = img.point(lambda x: 0 if x < 140 else 255)  # Simple thresholding
#     return img

# def normalize_text(text):
#     """Normalize text to make comparison easier."""
#     text = re.sub(r'\s+', ' ', text).strip()  # Normalize whitespace
#     text = re.sub(r'[^\w\s]', '', text).lower()  # Remove punctuation and convert to lowercase
#     return text

# def format_file_text(text, apply_spell_check=False):
#     """Format and clean extracted text."""
#     text = re.sub(r'\s+', ' ', text).strip()
#     text = re.sub(r'\.{3,}', '', text)
#     text = re.sub(r'\bcontrolled\b', '', text, flags=re.IGNORECASE)
#     # text = re.sub(r'[^a-zA-Z0-9\s.,?!]', '', text)  # Remove non-alphanumeric characters
#     try:
#         if detect(text) != 'en':
#             return ""
#     except Exception as e:
#         logging.error(f"Language detection failed: {e}")
#         return ""
#     return text.strip() #spell(text.strip()) if apply_spell_check else text.strip()

# def extract_page_text(page):
#     """Extract text and images from a single PDF page."""
#     try:
#         # Extract simple text
#         direct_text = page.get_text()
#         formatted_text = format_file_text(direct_text, apply_spell_check=False)

#         # Extract and process images for OCR
#         ocr_text = ""
#         images = page.get_images(full=True)
#         if images:
#             logging.info(f"Performing OCR on images in the page.")
#             pix = page.get_pixmap()
#             img_data = pix.tobytes("png")
#             img = Image.open(io.BytesIO(img_data))
#             img = preprocess_image(img)
#             ocr_text = pytesseract.image_to_string(img, lang='eng', config='--oem 3 --psm 6')
#             ocr_text = format_file_text(ocr_text, apply_spell_check=True)

#         # Avoid duplicate or repeated text
#         normalized_direct_text = normalize_text(formatted_text)
#         normalized_ocr_text = normalize_text(ocr_text)
#         if normalized_ocr_text not in normalized_direct_text:
#             combined_text = formatted_text + "\n" + ocr_text
#         else:
#             combined_text = formatted_text

#         return combined_text.strip()
#     except Exception as e:
#         logging.error(f"Error extracting text or OCR from page: {e}")
#         return ""

# def load_text_from_pdf(file_path):
#     """Load and extract text from a PDF file."""
#     doc = fitz.open(file_path)
#     text_pages = []
#     for page_num, page in enumerate(doc):
#         logging.info(f"Processing page {page_num + 1}.")
#         page_text = extract_page_text(page)
#         text_pages.append({
#             "page_number": page_num + 1,
#             "text": page_text,
#         })
#     return text_pages

# def load_text_from_docx(file_path):
#     """Load and extract text from a DOCX file."""
#     doc = Document(file_path)
#     text_pages = []
#     for page_num, para in enumerate(doc.paragraphs):
#         text = para.text.strip()
#         text_pages.append({
#             "page_number": page_num + 1,
#             "text": text,
#         })
#     return text_pages

# def load_text_from_txt(file_path):
#     """Load and extract text from a TXT file."""
#     with open(file_path, 'r', encoding='utf-8') as f:
#         text = f.read()
#     formatted_text = format_file_text(text)
#     return [{"page_number": 1, "text": formatted_text}]

# def load_text_from_file(file_path):
#     """Detect file type and load text accordingly."""
#     try:
#         file_path = file_path.lower()
#         if file_path.endswith('.pdf'):
#             return load_text_from_pdf(file_path)
#         elif file_path.endswith('.docx') or file_path.endswith('.doc'):
#             return load_text_from_docx(file_path)
#         elif file_path.endswith('.txt'):
#             return load_text_from_txt(file_path)
#         else:
#             logging.info(f"Unsupported file type: {file_path}")
#             return []
#     except Exception as e:
#         logging.error(f"Error loading file {file_path}: {e}", exc_info=True)
#         return []
